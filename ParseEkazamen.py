# -*- coding: utf-8 -*-
import re
import time
from docx import Document
import sqlite3
from pathlib import Path


times_start = time.time()
path = 'C:\PythonProj\mysite\db.sqlite3'
"""подключение к бд"""
db  = sqlite3.connect(path)
curs = db.cursor()

fio = re.compile('^[А-Я]{1}[а-я]*\s[А-Я]{1}[а-я]*\s[А-Я]{1}[а-я]*')


tim = re.compile('^\d{2}[:]\d{2}')
ddate = re.compile('^\d{2}[.]\d{2}[.]\d{2}')

def get_data_from_table(doc,stolbec):
    """
    функция предназначенная для получения списка строк извлеченных из ячеек указанного столюца таблицы, расположенной в
    .docx файле.

    :param doc: string - имя документа, содержащий таблицу с расписание экгзаменов, которая  парситься. Ввиду ограниечений модуля
    принимается только *.docx. Документ должен содержать только таблицу, без каких-либо заголовков и иных фрагментов текста.
    Таблица должна содержать 4 столбца следующего содержания:
        1) Предмет, с формой здачи и фио преподавателя
        2) Группы
        3) Дата и время
        4) Номер аудитории

    :param stolbec: int - номер столбца, из которого надо извлечь данные. Данные извлекаются ввиде строк, которые
    добавдяюстся в вовзвращаемый список.
    :return: list -  Список строк, содержавшихся во всех ячейках указанного столюца
    """
    document = Document(doc)
    tables = document.tables[0]
    data = []
    for row in tables.rows:
        data.append(row.cells[stolbec].paragraphs[0].text)
    return data

def search(pattern, massiv):
    """
    Функиция предназнаенная для поиска с помощью регулярных выражений в списке полученном в функции get_data_from_table().
    На выход функция получает список строк, на выходе - списко строк соответсвующих регулярному выражению
    :param pattern: Регулярные выражения используемые для поиска.
    :param massiv: Спискок внутри которого необходимо выполнить поиск
    :return: списко строк соответсвующих регулярному выражению в порядке их упоминиания в тексте.
    """
    finded = []
    for i in range(len(massiv)):
        search = re.findall(pattern, massiv[i])
        for j in range(len(search)):
            if search[j] is not None:
                finded.append(search[j])
    return finded

def unicum(massiv):
    """
    :param massiv:
    :return:
    """
    unic = []
    for i in range(len(massiv)):
        if massiv[i] not in unic:
            unic.append(massiv[i])
    return unic

def get_prep(doc):
    """
    :param doc:
    :return:
    """

    stolbec = 0
    prp = get_data_from_table(doc, stolbec)
    preps = search(fio, prp)
    preps_unic = unicum(preps)
    return preps_unic

def get_group(doc):
    """
    :return:
    """
    groups = re.compile('^\d{1}[А-Яа-я]*\d?')
    stolbec = 1
    grp = get_data_from_table(doc, stolbec)
    tmp_gr = search(groups, grp)
    group = unicum(tmp_gr)
    return group

def get_date(doc):
    """
    :return:
    """
    stolbec = 2
    tmp_date = get_data_from_table(doc, stolbec)
    srch_date = search(ddate,tmp_date)
    zdate = unicum(srch_date
    return zdate

def get_time(doc):
    """
    :doc
    :return:
    """
    stolbec = 2
    spl = []
    sltd = []
    sltd1 = []
    sltd2 = []
    pattern = ','

    tmp_time = get_data_from_table(doc, stolbec)
    for i in range(len(tmp_time)):
        sltd1 = re.split('\s', tmp_time[i])
        sltd.append(sltd1)
    sltd2 = search(tim,str(sltd))
    return sltd2



def get_auds(doc):
    stolbec =
    aud = re.compile('^\d*\D?')

    pass

doc = '2.docx'
func1 = get_group('2.docx')
print(func1)









